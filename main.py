
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import shutil, os, uuid, time
import openai
import pandas as pd
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

openai.api_key = os.getenv("OPENAI_API_KEY", "sk-...")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploaded"
RESULT_DIR = "translated"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(RESULT_DIR, exist_ok=True)

async def translate_text(text, source_lang="Chinese", target_lang="Vietnamese", model="gpt-4"):
    if not text.strip():
        return ""
    prompt = f"Dịch đoạn văn bản sau từ {source_lang} sang {target_lang} và trình bày song ngữ:\n\nNguyên văn:\n{text}\n\nBản dịch:"
    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return response.choices[0].message["content"].strip()
    except Exception as e:
        print(f"GPT error: {e}")
        return text

@app.post("/translate")
async def translate_file(
    file: UploadFile = File(...),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    format: str = Form("docx"),
):
    file_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")
    output_filename = f"translated_{file.filename.rsplit('.', 1)[0]}.docx"
    output_path = os.path.join(RESULT_DIR, output_filename)

    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    ext = file.filename.split(".")[-1].lower()
    start_time = time.time()

    try:
        progress, eta_seconds = 0, 0

        if ext in ["xlsx", "xls"]:
            df = pd.read_excel(input_path, header=None)
            total = df.shape[0] * df.shape[1]
            done = 0
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    original = str(df.iat[i, j])
                    tick = time.time()
                    translated = await translate_text(original, source_lang, target_lang)
                    df.iat[i, j] = f"{original}\n{translated}"
                    done += 1
                    elapsed = time.time() - tick
                    avg_time = (time.time() - start_time) / done
                    progress = int((done / total) * 100)
                    eta_seconds = int(avg_time * (total - done))
            output_path = os.path.join(RESULT_DIR, f"translated_{file.filename}")
            df.to_excel(output_path, index=False, header=False)

        elif ext == "docx":
            doc = Document(input_path)
            total = len(doc.paragraphs)
            for idx, para in enumerate(doc.paragraphs):
                tick = time.time()
                translated = await translate_text(para.text, source_lang, target_lang)
                para.text = f"{para.text}\n{translated}"
                elapsed = time.time() - tick
                avg_time = (time.time() - start_time) / (idx + 1)
                progress = int((idx + 1) / total * 100)
                eta_seconds = int(avg_time * (total - idx - 1))
            doc.save(output_path)

        elif ext == "pdf":
            doc = fitz.open(input_path)
            paragraphs = []
            for page in doc:
                text = page.get_text()
                paragraphs.extend(text.split("\n"))
            word_doc = Document()
            total = len(paragraphs)
            for idx, para in enumerate(paragraphs):
                tick = time.time()
                translated = await translate_text(para, source_lang, target_lang)
                word_doc.add_paragraph(f"{para}\n{translated}")
                elapsed = time.time() - tick
                avg_time = (time.time() - start_time) / (idx + 1)
                progress = int((idx + 1) / total * 100)
                eta_seconds = int(avg_time * (total - idx - 1))
            word_doc.save(output_path)

        elif ext in ["png", "jpg", "jpeg", "bmp"]:
            image = Image.open(input_path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            translated = await translate_text(text, source_lang, target_lang)
            word_doc = Document()
            word_doc.add_paragraph(f"{text}\n{translated}")
            word_doc.save(output_path)

        else:
            return JSONResponse({"error": "Định dạng file không được hỗ trợ."}, status_code=400)

        return {
            "translated_file_url": f"/download/{os.path.basename(output_path)}",
            "status": "done",
            "progress": 100,
            "eta_seconds": 0
        }

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/download/{filename}")
async def download_file(filename: str):
    path = os.path.join(RESULT_DIR, filename)
    if not os.path.exists(path):
        return JSONResponse({"error": "Không tìm thấy file."}, status_code=404)
    return FileResponse(path)
